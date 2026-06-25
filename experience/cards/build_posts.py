# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\e4user\Documents\ParticleworksProjects\pw-website\experience\cards"
REG = "https://particleworks-europe.com/experience/registration.php"
LOC = "📍 Modena, Italy 🇮🇹"
DATE = "📅 6-7 Oct 2026"
CTA = f"To attend Particleworks Experience 2026, register for free: {REG}"

posts = [
 dict(card="card-01-prometech.png", head="🚀 WHAT'S NEW IN PARTICLEWORKS 9.0 & GRANULEWORKS 4.0",
   intro="Issei Masaie and Iori Saigo from Prometech Software will unveil the latest Particleworks 9.0 and Granuleworks 4.0 releases — in two parts: what's new, then application examples and case studies.",
   lead="Discover what's new in particle-based simulation:",
   bullets=["Issei Masaie — What's New: major gains in simulation performance and GPU efficiency",
            "Extended multi-physics capabilities for complex fluid dynamics",
            "Iori Saigo — Application Examples & Case Studies from the latest developments"]),
 dict(card="card-02-mtu.png", head="🎯 SHOT PEENING SIMULATION",
   intro="Alpcan Güray from MTU Aero Engines AG will show how Particleworks models the shot peening of fatigue-critical aerospace components.",
   lead="Leveraging coupled CFD-DEM simulation, the team was able to:",
   bullets=["Accurately predict shot trajectories from machine, nozzle and air-pressure parameters",
            "Estimate surface coverage and the resulting residual-stress distribution",
            "Optimise robotic path planning for homogeneous peening"]),
 dict(card="card-03-unimore-stator.png", head="⚡ E-MOTOR STATOR COOLING",
   intro="Michelangelo Raimondo from the University of Modena and Reggio Emilia will present an integrated simulation approach for stator oil-jacket and jet cooling systems.",
   lead="Leveraging particle simulation, the study was able to achieve:",
   bullets=["A coupled MPS–FVM method that restores accuracy in critical wall-flow regions",
            "Reliable modelling of both wall-flows and free-flows",
            "Excellent agreement with experimental validation"]),
 dict(card="card-04-rdcfd-pump.png", head="🌡️ RECIPROCATING PUMP THERMAL ANALYSIS",
   intro="Leonardo Lanciotti from R&D CFD will show a conjugate-heat-transfer analysis of a reciprocating pump.",
   lead="Leveraging particle simulation, the team was able to:",
   bullets=["Identify hot spots and verify thermal limits at the piston seals",
            "Derive heat-transfer-coefficient distributions from an isothermal sloshing analysis",
            "Validate CHT results against experimental thermocouple measurements"]),
 dict(card="card-05-hesso-pelton.png", head="💧 PELTON TURBINE EROSION",
   intro="Jean Decaix from HES-SO Valais//Wallis will present Moving Particle Simulation of eroded Pelton runners.",
   lead="Leveraging GPU-accelerated MPS, the study was able to:",
   bullets=["Simulate real erosion geometries from new, mid-life and end-of-life runners",
            "Quantify torque losses of 6% and 6.6% for worn runners",
            "Capture complex free-surface, high-velocity flows on rotating buckets"]),
 dict(card="card-06-skf.png", head="⚙️ BEARING HEAT TRANSFER",
   intro="Lijun Cao and Mehul Pandya from SKF will show Particleworks capabilities for predicting heat transfer coefficients on bearing components.",
   lead="Leveraging particle simulation, SKF was able to achieve:",
   bullets=["Quick estimation of convective HTCs coupled with SKF's BEAST tool",
            "Validated benchmarks against literature correlations and ANSYS Fluent",
            "More accurate and robust bearing thermal analysis"]),
 dict(card="card-07-deepfluid.png", head="🫧 DIRECT AIR-IN-OIL MEASUREMENT",
   intro="Dr. Lukas Hafner from deepfluid will show how direct optical air-in-oil measurement enables more accurate simulation of gearings and hydraulic systems.",
   lead="Leveraging optical measurement and particle simulation, deepfluid was able to:",
   bullets=["Detect and quantify air bubbles in lubricants and hydraulic fluids",
            "Generate validated data to feed real air-in-oil behaviour into simulations",
            "Reduce uncertainties and the number of design-iteration steps"]),
 dict(card="card-08-trackone.png", head="🚜 UNDERCARRIAGE ROLLER LUBRICATION",
   intro="Leonardo Tiberi from Track One SRL will show a CFD 3D-MPS study of carrier-roller lubrication.",
   lead="Leveraging particle simulation, Track One was able to:",
   bullets=["Model the lubricant behaviour inside the roller layout under field conditions",
            "Investigate both fluid-dynamic and thermal aspects of the track/carrier roller",
            "Validate the model against bench-test roller temperatures with good correlation"]),
 dict(card="card-09-univance.png", head="⚙️ GEAR LUBRICATION & AIRFLOW",
   intro="Naohiro Fujita from Univance Corporation will show how Particleworks supports gear lubrication analysis and its expansion to R&D on airflow effects.",
   lead="Leveraging particle simulation, Univance was able to:",
   bullets=["Capture complex oil splashing and churning in the early design stage",
            "Visualise lubrication conditions and compare design options efficiently",
            "Extend the method to investigate airflow effects on gear lubrication"]),
 dict(card="card-10-iav.png", head="🫧 AIR BUBBLE DYNAMICS IN E-DRIVETRAINS",
   intro="René Kockisch from IAV GmbH will show air-bubble dynamics in the gear oil of an electric 3-speed commercial-vehicle drivetrain.",
   lead="Leveraging the IAV Particle Explorer and high-speed imaging, the team was able to:",
   bullets=["Quantify bubble formation, transport, coalescence and dissolution",
            "Reveal characteristic bubble size distributions and preferred transport paths",
            "Support optimisation of transmission design and drivetrain efficiency"]),
]

# announcement post (program release) — shown first, uses a program-link CTA
posts.insert(0, dict(card="card-00-program.png", head="📣 THE PROGRAM IS OUT",
   intro="The full programme for Particleworks Experience 2026 is online — 10 talks across two days in Modena, spanning e-drivetrains, e-motor cooling, bearings, gearboxes, gear lubrication, pumps, hydropower, aerospace and aeration.",
   lead="Join the meshless-CFD community on 6–7 October:",
   bullets=["10 validated industrial & academic case studies",
            "The release of Particleworks 9.0 and Granuleworks 4.0",
            "Hands-on workshops and networking at BPER FORUM Monzani"],
   imgs="card-00-program.png  →  card-00b-speakers.png  (2-slide carousel)",
   cta="See the full program and register for free: https://particleworks-europe.com/experience/program.php"))

# 30 companies to tag / target per post, keyed by card (matched to each talk's industry)
COMPANIES = {
 "card-00-program.png": "ZF, Bosch, Schaeffler, Continental, Vitesco Technologies, BorgWarner, Magna, Dana, GKN Automotive, Valeo, Mahle, Marelli, Stellantis, Volkswagen Group, BMW Group, Mercedes-Benz, Renault Group, Volvo Cars, Ferrari, Toyota, Hyundai Motor Group, Nidec, AVL, FEV, Ricardo, MTU Aero Engines, Rolls-Royce, Safran, SKF, Comer Industries",
 "card-01-prometech.png": "Robert Bosch, ZF Friedrichshafen, Schaeffler, Continental, Vitesco Technologies, BorgWarner, Magna, Valeo, Mahle, Dana, GKN Automotive, IAV, AVL, FEV, Ricardo, hofer powertrain, Bertrandt, EDAG Engineering, Porsche Engineering, Stellantis, BMW Group, Mercedes-Benz, Volkswagen Group, Toyota, Honda, Hyundai Motor Group, Nidec, Marelli, Denso, Aisin",
 "card-02-mtu.png": "Rolls-Royce, Safran, GE Aerospace, Pratt & Whitney, Honeywell Aerospace, Avio Aero, ITP Aero, GKN Aerospace, Collins Aerospace, Liebherr-Aerospace, Leonardo, Airbus, Boeing, Williams International, ArianeGroup, Curtiss-Wright Surface Technologies, Metal Improvement Company, Wheelabrator (Norican Group), Rösler Oberflächentechnik, Sintokogio, Sturm Maschinenbau, KSA Kugelstrahltechnik, Engineered Abrasives, Oerlikon Balzers, Bodycote, Doncasters Group, Howmet Aerospace, Rheinmetall, MTU Maintenance, Praxair Surface Technologies",
 "card-03-unimore-stator.png": "Vitesco Technologies, BorgWarner, ZF, Bosch, Mahle, Valeo, Nidec, Magna Powertrain, Dana TM4, Schaeffler, hofer powertrain, Equipmake, YASA, Marelli, Hyundai Mobis, Hyundai Transys, Aisin, Denso, Tesla, Rivian, Lucid Motors, BMW Group, Mercedes-Benz, Volkswagen Group, Porsche, Stellantis, BYD, Rimac Technology, Drive System Design, AVL",
 "card-04-rdcfd-pump.png": "KSB, Grundfos, Sulzer, Wilo, Bosch Rexroth, Parker Hannifin, Danfoss Power Solutions, Eaton, HAWE Hydraulik, Moog, Xylem, Flowserve, Weir Group, ITT, Pentair, Casappa, Bucher Hydraulics, Marzocchi Pompe, Kawasaki Precision Machinery, Bondioli & Pavesi, Interpump Group, HYDAC, Argo-Hytos, Graco, Pierburg, Continental, Caterpillar, Liebherr, Brevini, Duplomatic MS",
 "card-05-hesso-pelton.png": "Voith Hydro, Andritz Hydro, GE Vernova Hydro, Gilkes, Litostroj Power, ZECO Hydropower, Canyon Hydro, Global Hydro Energy, Mavel, Rainpower, EDF, Statkraft, Verbund, Enel Green Power, Hydro-Québec, Axpo, Alpiq, BKW, Iberdrola, Vattenfall, Norconsult, Stucky (Gruner Group), AFRY, WWS Wasserkraft, Toshiba Energy Systems, Sulzer, Tractebel, SN Power, Troy Hydro, Mecamidi",
 "card-06-skf.png": "Schaeffler, NSK, NTN, The Timken Company, JTEKT (Koyo), GMN Bearing, IKO Nippon Thompson, RBC Bearings, ZF, Bosch, Dana, GKN Automotive, BorgWarner, Comer Industries, Bonfiglioli, Carraro, Oerlikon Graziano, Allison Transmission, Eaton, Aisin, Voith, Flender, Siemens Gamesa, Vestas, Nordex, Liebherr, Caterpillar, Komatsu, Danfoss, Wittenstein",
 "card-07-deepfluid.png": "FUCHS, Klüber Lubrication, Shell, TotalEnergies, Castrol (BP), ExxonMobil, Petronas Lubricants, Liqui Moly, Bosch Rexroth, Parker Hannifin, Danfoss, HAWE Hydraulik, HYDAC, Moog, ZF, Schaeffler, Flender, SEW-Eurodrive, Bonfiglioli, Wittenstein, Bosch, Eaton, Comer Industries, Voith, GKN Automotive, Dana, Argo-Hytos, Stauff, MAHLE, Mann+Hummel",
 "card-08-trackone.png": "Caterpillar, Komatsu, Liebherr, John Deere, Hitachi Construction Machinery, Volvo CE, JCB, Develon (Doosan), Hyundai Construction Equipment, Kubota, CNH Industrial, Bobcat, Sany, XCMG, Zoomlion, Berco (thyssenkrupp), Italtractor ITM, Titan International, Camso (Michelin), USCO, Topy Industries, Prinoth, Kässbohrer (PistenBully), Claas, AGCO, Manitou, Terex, Wacker Neuson, Takeuchi, Yanmar",
 "card-09-univance.png": "Aisin, JTEKT, Hyundai Transys, ZF, GKN Automotive, Dana, BorgWarner, Schaeffler, Eaton, Allison Transmission, Oerlikon Graziano, Comer Industries, Bonfiglioli, Carraro, Brevini, Flender, SEW-Eurodrive, Wittenstein, Magna Powertrain, Marelli, Punch Powertrain, Vitesco Technologies, Drive System Design, Ricardo, AVL, FEV, Voith, Bharat Forge, Getrag, Bosch",
 "card-10-iav.png": "Daimler Truck, Volvo Trucks, Scania, MAN Truck & Bus, Traton Group, Iveco Group, PACCAR (DAF), CNH Industrial, ZF, Bosch, Vitesco Technologies, BorgWarner, Dana, Allison Transmission, Eaton, Voith, Schaeffler, Magna, AVL, FEV, Ricardo, Cummins, Deutz, Bharat Forge, Meritor, Hyundai Motor, BYD, Punch Powertrain, hofer powertrain, Drive System Design",
}

# ---------- DOCX ----------
doc = Document()
st = doc.styles["Normal"].font; st.name="Calibri"; st.size=Pt(11)
h = doc.add_heading("Particleworks Experience 2026 — LinkedIn Speaker Posts", level=0)
p = doc.add_paragraph("Ready-to-publish posts, one per speaker (styled like the 2024 edition). "
    "Swap the register link for a lnkd.in short link if preferred. Suggested image to attach is noted under each post.")
p.runs[0].italic=True
for i,po in enumerate(posts):
    doc.add_paragraph()
    hp=doc.add_paragraph(); r=hp.add_run(po["head"]); r.bold=True; r.font.size=Pt(13)
    doc.add_paragraph(po["intro"])
    doc.add_paragraph(po["lead"])
    for b in po["bullets"]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(LOC)
    doc.add_paragraph(DATE)
    doc.add_paragraph(po.get("cta", CTA))
    note=doc.add_paragraph(); rn=note.add_run(f"🖼️ Suggested image: {po.get('imgs', po['card'])}"); rn.italic=True; rn.font.color.rgb=RGBColor(0x6c,0x75,0x7d)
    cp=doc.add_paragraph(); rc=cp.add_run("🏷️ 30 companies to tag / target: "); rc.bold=True
    cp.add_run(COMPANIES[po["card"]])
    if i<len(posts)-1:
        sep=doc.add_paragraph("— — — — — — — — — — — — — — — — — — — —"); sep.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.save(os.path.join(OUT,"linkedin-posts-2026.docx"))

# ---------- Markdown ----------
md=["# Particleworks Experience 2026 — LinkedIn Speaker Posts","",
    f"Ready-to-publish posts, one per speaker (styled like the 2024 edition). Swap the register link for a `lnkd.in` short link if preferred.","",
    "Register link: " + REG,"","---",""]
for po in posts:
    md.append(f"## {po['head']}")
    md.append("")
    md.append(po["intro"]); md.append("")
    md.append(po["lead"])
    for b in po["bullets"]: md.append(f"- {b}")
    md.append("")
    md.append(LOC+"  "); md.append(DATE+"  "); md.append(po.get("cta", CTA))
    md.append("")
    md.append(f"*Suggested carousel: {po['imgs']}*" if po.get("imgs") else f"*Suggested image: `{po['card']}`*")
    md.append("")
    md.append("**🏷️ 30 companies to tag / target with this post:**  ")
    md.append(COMPANIES[po["card"]])
    md.append(""); md.append("---"); md.append("")
open(os.path.join(OUT,"linkedin-posts-2026.md"),"w",encoding="utf-8").write("\n".join(md))
print("wrote docx + md ->", OUT)
print("posts:", len(posts))
